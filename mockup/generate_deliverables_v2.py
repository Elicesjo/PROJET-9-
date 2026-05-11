"""
DWFA — Génération des livrables v2
- modele_print_DWFA_v2.docx
- mockup_DWFA_vue1.png / vue2.png / vue3.png
"""
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.lines import Line2D
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ─── Palette DWFA ───────────────────────────────────────────────────────────
C_DARK   = "#1A3A6B"
C_MID    = "#2368A0"
C_LIGHT  = "#5B9BD5"
C_XLIGHT = "#BDD7EE"
C_GREY   = "#F2F2F2"
C_GREY2  = "#D9D9D9"
C_WHITE  = "#FFFFFF"
C_WARN   = "#E67E22"
C_OK     = "#27AE60"

# ─── Helpers matplotlib ──────────────────────────────────────────────────────

def hex2rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) / 255 for i in (0, 2, 4))


def box(ax, x, y, w, h, fc=C_GREY, ec="#AAAAAA", lw=0.8, radius=0.012,
        label=None, fontsize=8, bold=False, fc_text=None, alpha=1.0):
    r = FancyBboxPatch((x, y), w, h,
                       boxstyle=f"round,pad={radius}",
                       linewidth=lw, edgecolor=ec,
                       facecolor=fc, alpha=alpha)
    ax.add_patch(r)
    if label:
        color = fc_text or ("#FFFFFF" if fc in (C_DARK, C_MID) else "#333333")
        ax.text(x + w / 2, y + h / 2, label,
                ha="center", va="center", fontsize=fontsize,
                fontweight="bold" if bold else "normal",
                color=color, linespacing=1.4)


def title_strip(ax, x, y, w, h, label):
    box(ax, x, y, w, h, fc=C_DARK, ec=C_DARK, label=label,
        fontsize=10, bold=True, fc_text=C_WHITE)


def kpi_card(ax, x, y, w, h, value, sublabel, color=C_MID):
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_LIGHT, lw=1.2)
    ax.text(x + w / 2, y + h * 0.62, value, ha="center", va="center",
            fontsize=12, fontweight="bold", color=color)
    ax.text(x + w / 2, y + h * 0.22, sublabel, ha="center", va="center",
            fontsize=6.5, color="#555555")


def filter_box(ax, x, y, w, h, label, kind="dropdown"):
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_LIGHT, lw=1.0)
    icon = "▾" if kind == "dropdown" else ("☑" if kind == "check" else "━")
    ax.text(x + 0.015, y + h / 2, icon, ha="left", va="center",
            fontsize=7, color=C_MID)
    ax.text(x + w / 2 + 0.01, y + h / 2, label, ha="center", va="center",
            fontsize=7, color="#333333")


def section_label(ax, x, y, label, fontsize=7.5, color="#888888"):
    ax.text(x, y, label, ha="left", va="center",
            fontsize=fontsize, color=color, style="italic")


# ── Mini graphiques réalistes ────────────────────────────────────────────────

def draw_bar_line(ax, x, y, w, h, title):
    """Bar chart + line overlay (accès eau mondiale)"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    # titre What-So What
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK, wrap=True)
    # zone de dessin
    mx, my = x + 0.025, y + 0.06
    mw, mh = w - 0.04, h - 0.10

    years = np.arange(2000, 2018)
    vals  = 55 + np.linspace(0, 25, len(years)) + np.random.default_rng(1).normal(0, 1, len(years))
    bar_w = mw / len(years) * 0.7
    for i, v in enumerate(vals):
        bx = mx + i * (mw / len(years))
        bh = (v / 100) * mh
        rect = plt.Rectangle((bx, my), bar_w, bh,
                              facecolor=C_LIGHT, edgecolor="none", alpha=0.85)
        ax.add_patch(rect)

    # line
    lx = [mx + i * (mw / len(years)) + bar_w / 2 for i in range(len(years))]
    ly = [my + (v / 100) * mh for v in vals]
    ax.plot(lx, ly, color=C_DARK, linewidth=1.5, zorder=5)
    ax.plot(lx[-1], ly[-1], "o", color=C_DARK, markersize=3, zorder=6)

    # axes
    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    ax.text(mx, my - 0.015, "2000", fontsize=5, color="#888888", ha="center")
    ax.text(mx + mw, my - 0.015, "2017", fontsize=5, color="#888888", ha="center")
    ax.text(mx - 0.01, my + mh / 2, "%", fontsize=5, color="#888888",
            va="center", ha="right")

    # légende
    patches = [mpatches.Patch(fc=C_LIGHT, label="% accès basique"),
               Line2D([0], [0], color=C_DARK, lw=1.5, label="Tendance")]
    ax.legend(handles=patches, loc="lower right",
              bbox_to_anchor=(x + w - 0.005, y + 0.01),
              fontsize=5, frameon=False)


def draw_line_pop(ax, x, y, w, h, title):
    """Line plot évolution population rurale vs totale"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK)
    mx, my = x + 0.03, y + 0.06
    mw, mh = w - 0.045, h - 0.10

    t = np.linspace(0, 1, 19)
    total = 6000 + t * 1500 + np.random.default_rng(2).normal(0, 30, 19)
    rural = 3200 - t * 400 + np.random.default_rng(3).normal(0, 25, 19)
    tmax  = total.max()

    lx = [mx + i * (mw / 18) for i in range(19)]
    ax.plot(lx, [my + (v / tmax) * mh for v in total],
            color=C_MID, linewidth=1.5, label="Population totale")
    ax.plot(lx, [my + (v / tmax) * mh for v in rural],
            color=C_XLIGHT, linewidth=1.5, linestyle="--", label="Population rurale")
    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    ax.text(mx, my - 0.015, "2000", fontsize=5, color="#888888", ha="center")
    ax.text(mx + mw, my - 0.015, "2018", fontsize=5, color="#888888", ha="center")
    ax.text(mx - 0.01, my + mh / 2, "Mrd", fontsize=5, color="#888888",
            va="center", ha="right")
    ax.legend(fontsize=5, loc="upper left",
              bbox_to_anchor=(x + 0.015, y + h - 0.04), frameon=False)


def draw_area_chart(ax, x, y, w, h, title):
    """Area chart stacked rural / urbain"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK)
    mx, my = x + 0.03, y + 0.06
    mw, mh = w - 0.045, h - 0.10

    t  = np.linspace(0, 1, 19)
    urban  = 40 + t * 20
    rural  = 60 - t * 20
    xs = np.array([mx + i * (mw / 18) for i in range(19)])
    yu = my + (urban / 100) * mh
    yr = my + (rural / 100) * mh

    ax.fill_between(xs, my, yr, color=C_XLIGHT, alpha=0.8, label="Rural")
    ax.fill_between(xs, yr, my + mh, color=C_MID, alpha=0.7, label="Urbain")
    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    ax.text(mx, my - 0.015, "2000", fontsize=5, color="#888888", ha="center")
    ax.text(mx + mw, my - 0.015, "2018", fontsize=5, color="#888888", ha="center")
    ax.legend(fontsize=5, loc="upper right",
              bbox_to_anchor=(x + w - 0.005, y + h - 0.04), frameon=False)


def draw_world_map(ax, x, y, w, h, title):
    """Carte monde simplifiée (choroplèthe simulée)"""
    box(ax, x, y, w, h, fc="#D6EAF8", ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK)

    # Continents simplifiés (ellipses)
    continents = [
        # (cx, cy, rx, ry, color, label)
        (x + w*0.25, y + h*0.60, w*0.09, h*0.17, C_MID,    "Afrique"),
        (x + w*0.22, y + h*0.78, w*0.06, h*0.10, C_LIGHT,  "Europe"),
        (x + w*0.60, y + h*0.65, w*0.14, h*0.20, C_XLIGHT, "Asie"),
        (x + w*0.17, y + h*0.38, w*0.07, h*0.14, C_LIGHT,  "Amériques"),
        (x + w*0.73, y + h*0.30, w*0.07, h*0.10, C_MID,    "Pacifique"),
    ]
    for cx, cy, rx, ry, col, lbl in continents:
        ell = mpatches.Ellipse((cx, cy), rx, ry,
                               facecolor=col, edgecolor="white", linewidth=0.7)
        ax.add_patch(ell)
        ax.text(cx, cy, lbl, ha="center", va="center",
                fontsize=4.5, color="white", fontweight="bold")

    # colorbar indicative
    for i, (col, txt) in enumerate([(C_XLIGHT, "Stable"), (C_LIGHT, "Modéré"),
                                     (C_MID, "Instable")]):
        bx = x + w * 0.05 + i * w * 0.12
        rect = plt.Rectangle((bx, y + 0.025), w * 0.10, 0.025,
                              facecolor=col, edgecolor="none")
        ax.add_patch(rect)
        ax.text(bx + w * 0.05, y + 0.018, txt,
                ha="center", va="top", fontsize=4.5, color="#555555")


def draw_grouped_bar(ax, x, y, w, h, title):
    """Grouped bar : accès basique vs safely managed par pays"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK)
    mx, my = x + 0.03, y + 0.07
    mw, mh = w - 0.045, h - 0.11

    pays = ["Pays A", "Pays B", "Pays C", "Pays D", "Pays E"]
    basic  = [85, 60, 45, 72, 30]
    safely = [40, 20, 10, 35, 5]
    n = len(pays)
    group_w = mw / n
    bw = group_w * 0.35

    for i, (b, s) in enumerate(zip(basic, safely)):
        gx = mx + i * group_w + group_w * 0.1
        bh_b = (b / 100) * mh
        bh_s = (s / 100) * mh
        rect_b = plt.Rectangle((gx, my), bw, bh_b, facecolor=C_LIGHT, edgecolor="none")
        rect_s = plt.Rectangle((gx + bw + 0.004, my), bw, bh_s,
                                facecolor=C_MID, edgecolor="none")
        ax.add_patch(rect_b)
        ax.add_patch(rect_s)
        ax.text(gx + bw, my - 0.015, pays[i], ha="center", va="top",
                fontsize=5, color="#555555")

    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    patches = [mpatches.Patch(fc=C_LIGHT, label="Basique"),
               mpatches.Patch(fc=C_MID, label="Safely managed")]
    ax.legend(handles=patches, fontsize=5,
              bbox_to_anchor=(x + w - 0.005, y + h - 0.04),
              loc="upper right", frameon=False)


def draw_scatter(ax, x, y, w, h, title, xlabel, ylabel, highlight=None):
    """Scatter plot générique"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.2, fontweight="bold",
            color=C_DARK)
    mx, my = x + 0.035, y + 0.07
    mw, mh = w - 0.05, h - 0.11

    rng = np.random.default_rng(42)
    xs  = rng.uniform(10, 90, 40)
    ys  = xs * 0.6 + rng.normal(0, 12, 40)
    ys  = np.clip(ys, 5, 95)

    colors = [C_MID] * 40
    if highlight:
        colors[5] = C_WARN
        colors[12] = C_OK

    for xi, yi, col in zip(xs, ys, colors):
        px = mx + (xi / 100) * mw
        py = my + (yi / 100) * mh
        ax.plot(px, py, "o", markersize=3.5, color=col, alpha=0.75)

    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    ax.plot([mx, mx], [my, my + mh], color="#AAAAAA", linewidth=0.5)
    ax.text(mx + mw / 2, my - 0.02, xlabel, ha="center", va="top",
            fontsize=5.5, color="#666666")
    ax.text(mx - 0.012, my + mh / 2, ylabel, ha="center", va="center",
            fontsize=5.5, color="#666666", rotation=90)

    if highlight:
        ax.plot([mx + 0.4 * mw, mx + 0.4 * mw], [my, my + mh],
                color="#DDDDDD", linewidth=0.8, linestyle="--")
        ax.plot([mx, mx + mw], [my + 0.5 * mh, my + 0.5 * mh],
                color="#DDDDDD", linewidth=0.8, linestyle="--")
        ax.text(mx + mw - 0.005, my + mh - 0.01, "Cible\nDWFA",
                ha="right", va="top", fontsize=4.5, color=C_OK)


def draw_stacked_bar(ax, x, y, w, h, title):
    """Percent stacked bar : pop urbaine vs rurale par pays"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK)
    mx, my = x + 0.02, y + 0.07
    mw, mh = w - 0.03, h - 0.11

    pays = ["A", "B", "C", "D", "E", "F"]
    urban_pct = [65, 30, 48, 72, 20, 55]
    n  = len(pays)
    bw = (mw / n) * 0.65

    for i, u in enumerate(urban_pct):
        gx = mx + i * (mw / n) + (mw / n) * 0.175
        h_u = (u / 100) * mh
        h_r = mh - h_u
        ax.add_patch(plt.Rectangle((gx, my + h_r), bw, h_u, facecolor=C_MID, edgecolor="none"))
        ax.add_patch(plt.Rectangle((gx, my), bw, h_r, facecolor=C_XLIGHT, edgecolor="none"))
        ax.text(gx + bw / 2, my - 0.012, pays[i], ha="center", va="top",
                fontsize=5, color="#555555")

    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    patches = [mpatches.Patch(fc=C_MID, label="Urbain"),
               mpatches.Patch(fc=C_XLIGHT, label="Rural")]
    ax.legend(handles=patches, fontsize=5,
              bbox_to_anchor=(x + w - 0.005, y + h - 0.04),
              loc="upper right", frameon=False)


def draw_line_water(ax, x, y, w, h, title):
    """Line plot accès eau pour un pays sélectionné"""
    box(ax, x, y, w, h, fc=C_WHITE, ec=C_GREY2, lw=1.0)
    ax.text(x + w / 2, y + h - 0.03, title,
            ha="center", va="top", fontsize=6.5, fontweight="bold",
            color=C_DARK)
    mx, my = x + 0.03, y + 0.065
    mw, mh = w - 0.045, h - 0.10

    t     = np.linspace(0, 1, 18)
    basic  = 45 + t * 30 + np.random.default_rng(7).normal(0, 1.5, 18)
    safely = 10 + t * 20 + np.random.default_rng(8).normal(0, 1.5, 18)
    xs    = np.array([mx + i * (mw / 17) for i in range(18)])

    ax.fill_between(xs, my, [my + (v / 100) * mh for v in basic],
                    color=C_XLIGHT, alpha=0.6)
    ax.plot(xs, [my + (v / 100) * mh for v in basic],
            color=C_LIGHT, linewidth=1.5, label="Basique")
    ax.plot(xs, [my + (v / 100) * mh for v in safely],
            color=C_MID, linewidth=1.5, linestyle="--", label="Safely managed")
    ax.plot([mx, mx + mw], [my, my], color="#AAAAAA", linewidth=0.5)
    ax.text(mx, my - 0.015, "2000", fontsize=5, color="#888888", ha="center")
    ax.text(mx + mw, my - 0.015, "2017", fontsize=5, color="#888888", ha="center")
    ax.legend(fontsize=5, bbox_to_anchor=(x + 0.015, y + h - 0.04),
              loc="upper left", frameon=False)


# ─── Filtres panel ───────────────────────────────────────────────────────────

def filters_panel(ax, x, y, filters):
    """Dessine un panel de filtres verticaux"""
    ax.text(x + 0.01, y + 0.02, "FILTRES", fontsize=6.5, fontweight="bold",
            color=C_DARK, va="bottom")
    for i, (lbl, kind) in enumerate(filters):
        fy = y - i * 0.115
        filter_box(ax, x, fy - 0.09, 0.185, 0.085, lbl, kind=kind)


# ─── Dashboard 1 — Mondiale ──────────────────────────────────────────────────

def build_vue1():
    fig, ax = plt.subplots(figsize=(18, 12))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("#F7FBFF")
    ax.set_facecolor("#F7FBFF")

    # Bandeau titre
    title_strip(ax, 0, 0.935, 1.0, 0.065,
                "VUE 1 — MONDIALE   |   Accès à l'eau potable dans le monde")

    # Sous-titre
    ax.text(0.5, 0.925, "Sélectionnez une année ou une région pour filtrer l'ensemble du tableau de bord",
            ha="center", va="top", fontsize=7.5, color="#555555", style="italic")

    # KPIs
    kpis = [
        ("74 %",  "Accès eau basique\n(monde, 2017)", C_MID),
        ("45 %",  "Accès safely\nmanaged (2017)", C_LIGHT),
        ("1,4 Mrd","Pop. sans accès\nà l'eau sûre", C_WARN),
        ("~829 k", "Décès WASH\nestimés (2016)", "#C0392B"),
    ]
    for i, (val, lbl, col) in enumerate(kpis):
        kpi_card(ax, 0.01 + i * 0.195, 0.855, 0.175, 0.065, val, lbl, color=col)

    # Ligne séparatrice
    ax.plot([0.01, 0.81], [0.848, 0.848], color=C_LIGHT, linewidth=0.5)

    # Graphiques haut
    draw_bar_line(ax, 0.01, 0.485, 0.375, 0.355,
                  "L'accès à l'eau progresse — mais 1,4 Mrd\nde personnes restent sans eau sûre en 2017")
    draw_line_pop(ax, 0.40, 0.485, 0.375, 0.355,
                  "La population mondiale croît — la part\nurbaine dépasse la part rurale depuis 2007")

    # Graphiques bas
    draw_area_chart(ax, 0.01, 0.085, 0.375, 0.380,
                    "La population urbaine dépasse la rurale —\nla transition urbaine s'accélère post-2010")
    draw_world_map(ax, 0.40, 0.085, 0.375, 0.380,
                   "Instabilité politique concentrée en Afrique\net Moyen-Orient — zones prioritaires DWFA")

    # Panel filtres
    ax.text(0.82, 0.91, "FILTRES", fontsize=7.5, fontweight="bold",
            color=C_DARK, va="top")
    ax.plot([0.815, 0.815], [0.08, 0.93], color=C_LIGHT, linewidth=0.8, alpha=0.5)

    filt = [("Année (2000–2017)", "dropdown"),
            ("Région OMS", "dropdown"),
            ("Métrique principale", "dropdown"),
            ("Seuil stabilité pol.", "slider")]
    for i, (lbl, kind) in enumerate(filt):
        filter_box(ax, 0.82, 0.86 - i * 0.11, 0.17, 0.085, lbl, kind=kind)

    box(ax, 0.82, 0.30, 0.17, 0.22, fc=C_WHITE, ec=C_LIGHT)
    ax.text(0.905, 0.51, "Exclure pays\nsans données", ha="center",
            va="top", fontsize=6.5, color="#444444")
    for i, reg in enumerate(["Afrique", "Amériques", "Europe", "Asie SE"]):
        ry = 0.47 - i * 0.04
        ax.text(0.838, ry, "☐", fontsize=7, color=C_MID, va="center")
        ax.text(0.855, ry, reg, fontsize=6, color="#444444", va="center")

    # Légende palette
    ax.text(0.82, 0.26, "Stabilité politique :", fontsize=6, color="#555555")
    for i, (col, lbl) in enumerate([(C_XLIGHT, "Stable"), (C_LIGHT, "Modéré"),
                                     (C_MID, "Instable")]):
        rect = plt.Rectangle((0.82 + i * 0.058, 0.23), 0.05, 0.02,
                              facecolor=col, edgecolor="none")
        ax.add_patch(rect)
        ax.text(0.845 + i * 0.058, 0.222, lbl, ha="center",
                fontsize=5.5, color="#555555")

    ax.text(0.5, 0.04,
            "Source : OMS (BasicAndSafelyManagedDrinkingWaterServices, MortalityRateAttributedToWater) · FAO (Population) · Banque Mondiale (PoliticalStability)",
            ha="center", va="bottom", fontsize=5.5, color="#AAAAAA")
    ax.text(0.01, 0.04, "DWFA — Mission eau potable | Vue Mondiale",
            ha="left", va="bottom", fontsize=5.5, color="#AAAAAA")

    plt.tight_layout(pad=0.3)
    plt.savefig("mockup_DWFA_vue1.png", dpi=150, bbox_inches="tight",
                facecolor="#F7FBFF")
    print("mockup_DWFA_vue1.png généré.")
    plt.close()


# ─── Dashboard 2 — Continentale ──────────────────────────────────────────────

def build_vue2():
    fig, ax = plt.subplots(figsize=(18, 12))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("#F7FBFF")

    title_strip(ax, 0, 0.935, 1.0, 0.065,
                "VUE 2 — CONTINENTALE   |   Continent sélectionné : Afrique")
    ax.text(0.5, 0.925,
            "Sélectionnez un continent — les indicateurs s'actualisent automatiquement",
            ha="center", va="top", fontsize=7.5, color="#555555", style="italic")

    kpis = [
        ("58 %", "Accès eau\nbasique moy.", C_MID),
        ("12 %", "Accès safely\nmanaged moy.", C_LIGHT),
        ("54",   "Pays analysés\nsur le continent", C_DARK),
        ("-0.8", "Stabilité pol.\nmoyenne", C_WARN),
    ]
    for i, (val, lbl, col) in enumerate(kpis):
        kpi_card(ax, 0.01 + i * 0.195, 0.855, 0.175, 0.065, val, lbl, color=col)

    ax.plot([0.01, 0.81], [0.848, 0.848], color=C_LIGHT, linewidth=0.5)

    # Grouped bar (large, en haut)
    draw_grouped_bar(ax, 0.01, 0.485, 0.775, 0.350,
                     "Les pays d'Afrique sub-saharienne cumulent faible accès\nbasique ET quasi-absence de services de qualité")

    # Bas gauche : stacked bar
    draw_stacked_bar(ax, 0.01, 0.090, 0.370, 0.375,
                     "Les pays les moins urbanisés sont aussi\nceux avec le plus faible accès à l'eau")

    # Bas droite : scatter Domaine 3
    draw_scatter(ax, 0.41, 0.090, 0.375, 0.375,
                 "Domaine 3 — Consulting : pays avec\nbonne politique ET stabilité suffisante",
                 "Stabilité politique →", "Efficacité politique eau →",
                 highlight=True)

    # Panel filtres
    ax.text(0.82, 0.91, "FILTRES", fontsize=7.5, fontweight="bold", color=C_DARK)
    ax.plot([0.815, 0.815], [0.08, 0.93], color=C_LIGHT, linewidth=0.8, alpha=0.5)

    filt = [("Continent", "dropdown"),
            ("Année", "dropdown"),
            ("Seuil stabilité pol.", "slider"),
            ("Pays (multi-sélection)", "check")]
    for i, (lbl, kind) in enumerate(filt):
        filter_box(ax, 0.82, 0.86 - i * 0.11, 0.17, 0.085, lbl, kind=kind)

    # Note limites
    box(ax, 0.82, 0.15, 0.17, 0.14, fc="#FEF9E7", ec=C_WARN, lw=0.8)
    ax.text(0.905, 0.285, "⚠ Limites données", ha="center",
            fontsize=6, fontweight="bold", color=C_WARN, va="top")
    ax.text(0.905, 0.258,
            "Safely managed :\nnombreuses valeurs\nmanquantes\n\nMortalité :\n2016 uniquement",
            ha="center", fontsize=5.5, color="#555555", va="top", linespacing=1.5)

    ax.text(0.5, 0.04,
            "Source : OMS · FAO · Banque Mondiale",
            ha="center", va="bottom", fontsize=5.5, color="#AAAAAA")
    ax.text(0.01, 0.04, "DWFA — Mission eau potable | Vue Continentale",
            ha="left", va="bottom", fontsize=5.5, color="#AAAAAA")

    plt.tight_layout(pad=0.3)
    plt.savefig("mockup_DWFA_vue2.png", dpi=150, bbox_inches="tight",
                facecolor="#F7FBFF")
    print("mockup_DWFA_vue2.png généré.")
    plt.close()


# ─── Dashboard 3 — Nationale ─────────────────────────────────────────────────

def build_vue3():
    fig, ax = plt.subplots(figsize=(18, 12))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("#F7FBFF")

    title_strip(ax, 0, 0.935, 1.0, 0.065,
                "VUE 3 — NATIONALE   |   Pays sélectionné : [Sénégal]")
    ax.text(0.5, 0.925,
            "Sélectionnez un pays — diagnostic des 3 domaines d'expertise DWFA",
            ha="center", va="top", fontsize=7.5, color="#555555", style="italic")

    kpis = [
        ("67 %",  "Accès eau\nbasique (2017)", C_MID),
        ("10 %",  "Accès safely\nmanaged (2017)", C_LIGHT),
        ("~6 200","Décès WASH\nestimés (2016)", "#C0392B"),
        ("-0.4",  "Stabilité\npolitique (2017)", C_WARN),
    ]
    for i, (val, lbl, col) in enumerate(kpis):
        kpi_card(ax, 0.01 + i * 0.195, 0.855, 0.175, 0.065, val, lbl, color=col)

    ax.plot([0.01, 0.81], [0.848, 0.848], color=C_LIGHT, linewidth=0.5)

    # Line plot population (large en haut)
    draw_line_pop(ax, 0.01, 0.490, 0.375, 0.345,
                  "Population en forte croissance — la part\nrurale se réduit mais reste majoritaire")

    # Line plot accès eau
    draw_line_water(ax, 0.41, 0.490, 0.375, 0.345,
                    "L'accès basique progresse — l'écart avec\nle safely managed reste structurel")

    # Scatter Domaine 1
    draw_scatter(ax, 0.01, 0.090, 0.370, 0.380,
                 "Domaine 1 — Création : faible accès\net population dispersée → priorité rurale",
                 "% population urbaine →", "% accès eau →")

    # Scatter Domaine 2
    draw_scatter(ax, 0.41, 0.090, 0.370, 0.380,
                 "Domaine 2 — Modernisation : accès basique\ncorrect mais safely managed insuffisant",
                 "% accès basique →", "% accès safely managed →")

    # Panel filtres
    ax.text(0.82, 0.91, "FILTRES", fontsize=7.5, fontweight="bold", color=C_DARK)
    ax.plot([0.815, 0.815], [0.08, 0.93], color=C_LIGHT, linewidth=0.8, alpha=0.5)

    filt = [("Pays", "dropdown"),
            ("Continent", "dropdown"),
            ("Année (plage)", "slider"),
            ("Granularité", "dropdown")]
    for i, (lbl, kind) in enumerate(filt):
        filter_box(ax, 0.82, 0.86 - i * 0.11, 0.17, 0.085, lbl, kind=kind)

    # Diagnostic DWFA
    box(ax, 0.82, 0.12, 0.17, 0.28, fc=C_WHITE, ec=C_MID, lw=1.0)
    ax.text(0.905, 0.395, "Diagnostic DWFA", ha="center",
            fontsize=6.5, fontweight="bold", color=C_DARK, va="top")
    diag = [
        (C_OK,   "D1 Création",      "Pertinent"),
        (C_OK,   "D2 Modernisation", "Pertinent"),
        (C_WARN, "D3 Consulting",    "Instabilité"),
    ]
    for i, (col, dom, stat) in enumerate(diag):
        dy = 0.355 - i * 0.07
        circ = plt.Circle((0.838, dy), 0.008, color=col)
        ax.add_patch(circ)
        ax.text(0.852, dy, f"{dom} : {stat}", va="center",
                fontsize=5.5, color="#333333")

    ax.text(0.5, 0.04,
            "Source : OMS · FAO · Banque Mondiale",
            ha="center", va="bottom", fontsize=5.5, color="#AAAAAA")
    ax.text(0.01, 0.04, "DWFA — Mission eau potable | Vue Nationale",
            ha="left", va="bottom", fontsize=5.5, color="#AAAAAA")

    plt.tight_layout(pad=0.3)
    plt.savefig("mockup_DWFA_vue3.png", dpi=150, bbox_inches="tight",
                facecolor="#F7FBFF")
    print("mockup_DWFA_vue3.png généré.")
    plt.close()


# ─── Modèle print Word ───────────────────────────────────────────────────────

ROWS = [
    ("Vue d'ensemble mondiale — accès eau potable",
     "% accès safely managed (agrégation mondiale — MOYENNE pondérée)",
     "KPI + Bar + Line",
     "Mondiale",
     "BasicAndSafely…csv",
     "✅ 2000–2017 / ⚠ safely managed : nombreuses valeurs manquantes"),
    ("Évolution population mondiale (rurale vs totale)",
     "Population totale, population rurale (milliers × 1 000)",
     "Line plot (2 séries)",
     "Mondiale",
     "Population.csv",
     "✅ 2000–2018"),
    ("Répartition urbain / rural dans le temps",
     "% pop. urbaine = pop_urban / pop_total × 100 (champ calculé)",
     "Area chart / Stacked bar",
     "Mondiale & Continentale",
     "Population.csv",
     "✅ 2000–2018"),
    ("Stabilité politique — identifier zones instables",
     "Moyenne indice stabilité politique par continent (MOYENNE)",
     "Carte choroplèthe",
     "Mondiale",
     "PoliticalStability.csv",
     "⚠ pas annuel — pas irréguliers (2000, 2002…)"),
    ("Comparaison accès eau entre pays d'un continent",
     "% accès basique, % accès safely managed",
     "Grouped bar plot",
     "Continentale",
     "BasicAndSafely…csv",
     "✅ 2000–2017 / ⚠ safely managed lacunaire"),
    ("Répartition urbain / rural par pays",
     "% pop. urbaine (champ calculé)",
     "Percent stacked bar",
     "Continentale",
     "Population.csv",
     "✅ 2000–2018"),
    ("Domaine 1 — Création : cibler selon densité urbaine/rurale",
     "% accès eau potable × % pop. urbaine (2 axes scatter)",
     "Scatter plot",
     "Nationale",
     "BasicAndSafely…csv + Population.csv",
     "✅ jointure sur Country + Year"),
    ("Domaine 2 — Modernisation : services à améliorer",
     "% accès basique × % accès safely managed",
     "Scatter plot",
     "Nationale & Continentale",
     "BasicAndSafely…csv",
     "⚠ safely managed souvent vide"),
    ("Domaine 3 — Consulting : politique efficace + contexte stable",
     "Stabilité politique × Efficacité pol. eau (mortalité + accès)",
     "Scatter plot",
     "Continentale",
     "PoliticalStability.csv + Mortality…csv + BasicAndSafely…csv",
     "⚠ mortalité 2016 uniquement"),
    ("Nb décès WASH par pays",
     "Nb décès = taux_mortalité × pop_total / 100 000 (champ calculé — colonne WASH deaths vide)",
     "KPI + Bar chart",
     "Toutes vues",
     "Mortality…csv + Population.csv",
     "⚠ 2016 uniquement"),
    ("Évolution accès eau dans le temps pour un pays",
     "% accès basique, % accès safely managed (2000–2017)",
     "Line plot (2 séries)",
     "Nationale",
     "BasicAndSafely…csv",
     "✅ 2000–2017"),
]

FILTERS = [
    ("Année",               "Temporel",          "Slider ou liste déroulante", "Toutes vues"),
    ("Région OMS / Continent","Géographique — qualitatif","Liste déroulante",  "Mondiale & Continentale"),
    ("Pays",                "Géographique — qualitatif","Liste déroulante",    "Nationale"),
    ("Seuil stabilité pol.","Quantitatif",       "Curseur paramétrable",       "Continentale"),
]

CALC_FIELDS = [
    ("% Population urbaine",         "pop_urban / pop_total × 100",                   "Domaine 1 — scatter"),
    ("Nb décès WASH estimés",        "taux_mortalité × pop_total / 100 000",           "KPI toutes vues"),
    ("Stabilité moy. continentale",  "AVG(political_stability) GROUP BY région",       "Carte monde"),
    ("Efficacité politique eau",     "Norm(1 – taux_mortalité) + Norm(% safely managed)","Domaine 3 — scatter"),
    ("Population (en personnes)",    "Population_csv × 1 000",                        "Conversion unité"),
]

DATA_STATUS = [
    ("% accès eau basique",    "BasicAndSafely…csv",        "✅ 2000–2017, Total/Urban/Rural"),
    ("% accès safely managed", "BasicAndSafely…csv",        "⚠ nombreuses valeurs manquantes"),
    ("Taux mortalité WASH",    "Mortality…csv",              "⚠ 2016 uniquement"),
    ("Nb décès WASH bruts",    "Mortality…csv",              "❌ colonne vide → à calculer"),
    ("Population",             "Population.csv",             "✅ 2000–2018, Total/Urban/Rural/M/F"),
    ("% pop. urbaine",         "Population.csv",             "✅ à calculer"),
    ("Stabilité politique",    "PoliticalStability.csv",     "⚠ pas irréguliers, pas tous les ans"),
    ("Région / Continent",     "RegionCountry.csv",          "✅ 194 pays, 6 régions OMS"),
]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def add_header_row(table, headers, bg="1A3A6B"):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def fill_row(row, values, alt=False):
    bg = "E8F1FB" if alt else "FFFFFF"
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].font.size = Pt(8.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def h2_style(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x23, 0x68, 0xA0)
    return p


def build_model_print():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Page de garde ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run("\n\n")

    t = doc.add_heading("DWFA — Drinking Water For All", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
    t.runs[0].font.size = Pt(20)

    t2 = doc.add_heading("Modèle Print — Tableau de bord accès à l'eau", level=2)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.color.rgb = RGBColor(0x23, 0x68, 0xA0)

    sub = doc.add_paragraph(
        "Document de synthèse des indicateurs retenus · Version 1.0 · Avril 2026"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].italic = True
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    intro = doc.add_paragraph(
        "Ce document synthétise les choix d'indicateurs, de visualisations et de filtres "
        "pour les 3 vues du tableau de bord DWFA. Il constitue le livrable préalable à la "
        "construction du dashboard dans Tableau Public."
    )
    intro.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── Section 1 : Indicateurs ───────────────────────────────────────────
    h2_style(doc, "1. Indicateurs et visualisations")

    headers = ["Besoin utilisateur", "Mesures", "Visualisation", "Vue",
               "Source (fichier)", "Disponibilité / Limites"]
    col_w = [Cm(4.5), Cm(4.2), Cm(2.8), Cm(2.0), Cm(3.2), Cm(4.2)]

    table = doc.add_table(rows=1 + len(ROWS), cols=6)
    table.style = "Table Grid"
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_w[i]
    add_header_row(table, headers)
    for idx, row_data in enumerate(ROWS):
        fill_row(table.rows[idx + 1], row_data, alt=(idx % 2 == 0))

    doc.add_paragraph()

    # ── Section 2 : Filtres ───────────────────────────────────────────────
    h2_style(doc, "2. Filtres interactifs")

    f_table = doc.add_table(rows=1 + len(FILTERS), cols=4)
    f_table.style = "Table Grid"
    add_header_row(f_table, ["Variable filtrée", "Type", "Implémentation", "Vue(s)"])
    for idx, row_data in enumerate(FILTERS):
        fill_row(f_table.rows[idx + 1], row_data, alt=(idx % 2 == 0))

    doc.add_paragraph()

    # ── Section 3 : Champs calculés ───────────────────────────────────────
    h2_style(doc, "3. Champs calculés")

    c_table = doc.add_table(rows=1 + len(CALC_FIELDS), cols=3)
    c_table.style = "Table Grid"
    add_header_row(c_table, ["Indicateur", "Formule", "Utilité"])
    for idx, row_data in enumerate(CALC_FIELDS):
        fill_row(c_table.rows[idx + 1], row_data, alt=(idx % 2 == 0))

    doc.add_paragraph()

    # ── Section 4 : État des données ─────────────────────────────────────
    h2_style(doc, "4. État des données disponibles")

    d_table = doc.add_table(rows=1 + len(DATA_STATUS), cols=3)
    d_table.style = "Table Grid"
    add_header_row(d_table, ["Indicateur", "Fichier source", "Disponibilité"])
    for idx, row_data in enumerate(DATA_STATUS):
        fill_row(d_table.rows[idx + 1], row_data, alt=(idx % 2 == 0))

    doc.add_paragraph()

    # ── Section 5 : Points d'attention ───────────────────────────────────
    h2_style(doc, "5. Points d'attention")
    limits = [
        "Mortalité WASH disponible pour 2016 uniquement — les comparaisons temporelles sur cet indicateur sont impossibles.",
        "La colonne 'WASH deaths' dans le fichier Mortality est vide — le nombre de décès doit être calculé via : taux × population / 100 000.",
        "L'indicateur 'safely managed' présente de nombreuses valeurs manquantes, notamment pour les pays les moins avancés — à mentionner dans la présentation.",
        "La stabilité politique n'est pas disponible pour toutes les années (pas irréguliers) — interpolation ou dernière valeur connue à envisager.",
        "La population est exprimée en milliers dans Population.csv → multiplier par 1 000 dans tous les champs calculés.",
        "Filtre stabilité politique : l'indice peut être négatif (instabilité forte) → le seuil doit être paramétrable par l'utilisateur.",
        "Palette bleue obligatoire sur l'ensemble du dashboard. Accessibilité : contrastes suffisants, pas de rouge/vert seuls (daltonisme).",
    ]
    for lim in limits:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(lim).font.size = Pt(9)

    doc.save("modele_print_DWFA_v2.docx")
    print("modele_print_DWFA_v2.docx généré.")


if __name__ == "__main__":
    np.random.seed(0)
    build_model_print()
    build_vue1()
    build_vue2()
    build_vue3()
